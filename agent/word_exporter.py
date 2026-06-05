"""
Word and PDF document exporter for job listings.
Creates Word files with job information grouped by match percentage.
Each score range gets its own Word file + a separate PDF CV tailored for that range.
"""

import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Dict, Tuple

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .models import Job


class ScoredJob:
    """Job with AI scoring result."""
    def __init__(self, job: Job, match_score: int, cover_letter: str = ""):
        self.job = job
        self.match_score = match_score
        self.cover_letter = cover_letter
    
    @property
    def score_range(self) -> str:
        """Get the score range label."""
        if self.match_score >= 96:
            return "96-100"
        elif self.match_score >= 91:
            return "91-95"
        elif self.match_score >= 86:
            return "86-90"
        elif self.match_score >= 80:
            return "80-85"
        else:
            return None  # Below 80% - skip


class WordExporter:
    """Export job listings to Word documents, grouped by score range."""
    
    SCORE_RANGES = [
        ("96-100", "jobs_96_100.docx"),
        ("91-95", "jobs_91_95.docx"),
        ("86-90", "jobs_86_90.docx"),
        ("80-85", "jobs_80_85.docx"),
    ]
    
    def __init__(self, output_dir: str = "."):
        self.output_dir = Path(output_dir)
        self._ensure_dir()
    
    def _ensure_dir(self):
        """Ensure the directory exists."""
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def _add_hyperlink(self, paragraph, url: str, text: str):
        """Add a hyperlink to a paragraph."""
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '0066CC')
        rPr.append(color)
        
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        
        new_run.append(rPr)
        
        text_elem = OxmlElement('w:t')
        text_elem.text = text
        new_run.append(text_elem)
        
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
    
    def _style_header_row(self, row):
        """Style header row with blue background and white text."""
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '0066CC')
            cell._tc.get_or_add_tcPr().append(shading)
    
    def export_jobs_with_scores(self, scored_jobs: List[ScoredJob], profile_name: str = "",
                                 cv_texts: Optional[Dict[str, str]] = None) -> List[Tuple[str, str, int, str]]:
        """
        Export jobs to separate Word files based on score ranges.
        Each file contains jobs in that range + a common CV for that range.
        
        Args:
            scored_jobs: List of ScoredJob objects
            profile_name: Candidate name
            cv_texts: Dict mapping score_range -> CV text string
            
        Returns:
            List of (range_label, word_filepath, job_count, pdf_filepath) tuples
        """
        # Group jobs by score range
        grouped = {range_label: [] for range_label, _ in self.SCORE_RANGES}
        
        for scored_job in scored_jobs:
            if scored_job.score_range and scored_job.score_range in grouped:
                grouped[scored_job.score_range].append(scored_job)
        
        if cv_texts is None:
            cv_texts = {}
        
        # Create Word files + PDF CVs for each range
        created_files = []
        for range_label, filename in self.SCORE_RANGES:
            jobs = grouped[range_label]
            if not jobs:
                continue
            
            word_filepath = self.output_dir / filename
            cv_text = cv_texts.get(range_label, "")
            
            # Create Word document
            self._create_range_document(word_filepath, jobs, range_label, cv_text, profile_name)
            
            # Create PDF CV
            pdf_filename = f"cv_{range_label.replace('-', '_')}.pdf"
            pdf_filepath = self.output_dir / pdf_filename
            self._create_cv_pdf(pdf_filepath, cv_text, profile_name, range_label)
            
            created_files.append((range_label, str(word_filepath), len(jobs), str(pdf_filepath)))
        
        return created_files
    
    @staticmethod
    def _sanitize_for_pdf(text: str) -> str:
        """Replace Unicode characters that fpdf2's default font can't handle."""
        replacements = {
            '\u2014': '--',   # em dash
            '\u2013': '-',    # en dash
            '\u2018': "'",   # left single quote
            '\u2019': "'",   # right single quote
            '\u201c': '"',   # left double quote
            '\u201d': '"',   # right double quote
            '\u2022': '-',   # bullet
            '\u2026': '...',  # ellipsis
            '\u00a0': ' ',    # non-breaking space
            '\u2023': '-',   # triangular bullet
            '\u25cf': '*',   # filled circle
            '\u25cb': 'o',   # open circle
            '\u2713': '[OK]', # checkmark
            '\u2717': '[X]',  # cross
            '\u2192': '->',  # arrow
            '\u2605': '*',   # star
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        # Fallback: encode to latin-1, replacing anything else
        text = text.encode('latin-1', errors='replace').decode('latin-1')
        return text
    
    def _create_cv_pdf(self, filepath: Path, cv_text: str, profile_name: str, range_label: str):
        """Create a PDF file from CV text using fpdf2."""
        try:
            from fpdf import FPDF
        except ImportError:
            import logging
            logging.getLogger(__name__).warning("fpdf2 not installed - skipping PDF generation")
            return
        
        if not cv_text:
            return
        
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()
        
        # Title
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 10, txt=self._sanitize_for_pdf(profile_name or "CV"), ln=True, align="C")
        pdf.set_font("Helvetica", "I", 10)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 8, txt=self._sanitize_for_pdf(f"Tailored for {range_label}% Match Jobs -- Generated {datetime.now().strftime('%Y-%m-%d')}"), ln=True, align="C")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(6)
        
        # Parse and render CV text
        lines = cv_text.split('\n')
        in_bullet = False
        
        for line in lines:
            stripped = line.strip()
            
            if not stripped:
                pdf.ln(3)
                continue
            
            # Detect section headers (ALL CAPS or ends with colon and is short)
            is_section_header = (
                (stripped.isupper() and len(stripped) > 3) or
                (stripped.endswith(':') and len(stripped) < 60 and not stripped[0].islower()) or
                re.match(r'^#{1,3}\s+', stripped)
            )
            
            # Detect bullet points
            is_bullet = stripped.startswith(('•', '-', '·', '*')) or re.match(r'^\d+[\.\)]\s+', stripped)
            
            # Clean markdown headers
            clean_text = re.sub(r'^#{1,3}\s+', '', stripped)
            clean_text = clean_text.replace('**', '').replace('__', '')
            clean_text = self._sanitize_for_pdf(clean_text)
            
            if is_section_header:
                pdf.ln(4)
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(0, 51, 153)
                pdf.cell(0, 8, txt=clean_text, ln=True)
                pdf.set_draw_color(0, 51, 153)
                pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
                pdf.ln(3)
                pdf.set_text_color(0, 0, 0)
                in_bullet = False
                
            elif is_bullet:
                bullet_text = re.sub(r'^[\s]*[\-\•\·\*]\s*', '', clean_text)
                bullet_text = re.sub(r'^\d+[\.\)]\s*', '', bullet_text)
                pdf.set_font("Helvetica", "", 10)
                pdf.set_x(pdf.get_x() + 5)
                pdf.cell(5, 6, txt="-")
                pdf.multi_cell(175, 6, txt=bullet_text)
                in_bullet = True
                
            else:
                pdf.set_font("Helvetica", "", 10)
                pdf.multi_cell(190, 6, txt=clean_text)
                in_bullet = False
        
        pdf.output(str(filepath))
    
    def _create_range_document(self, filepath: Path, scored_jobs: List[ScoredJob],
                               range_label: str, cv_text: str, profile_name: str):
        """Create a Word document for a specific score range with jobs table + embedded CV."""
        doc = Document()
        
        # Title
        title = doc.add_heading(f'Jobs Matching {range_label}%', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        if profile_name:
            profile_para = doc.add_paragraph()
            profile_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = profile_para.add_run(f"Candidate: {profile_name}")
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        
        # Summary
        summary = doc.add_paragraph()
        run = summary.add_run(f"Total Jobs in {range_label}% Range: {len(scored_jobs)}")
        run.font.bold = True
        run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Jobs table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Job Role"
        header_cells[1].text = "Company"
        header_cells[2].text = "Match %"
        header_cells[3].text = "Link"
        self._style_header_row(table.rows[0])
        
        # Add job rows
        for scored_job in scored_jobs:
            row = table.add_row()
            row.cells[0].text = scored_job.job.title
            row.cells[1].text = scored_job.job.company
            row.cells[2].text = f"{scored_job.match_score}%"
            
            link_para = row.cells[3].paragraphs[0]
            if scored_job.job.url:
                self._add_hyperlink(link_para, scored_job.job.url, "Apply Link")
            else:
                link_para.text = "N/A"
            
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # CV Section
        if cv_text:
            cv_heading = doc.add_heading('Tailored CV for This Score Range', level=1)
            cv_heading.style.font.color.rgb = RGBColor(0, 102, 204)
            
            cv_note = doc.add_paragraph()
            cv_note_run = cv_note.add_run(
                "This CV has been AI-generated and tailored for jobs in this match percentage range. "
                "A separate PDF version is also available."
            )
            cv_note_run.font.size = Pt(9)
            cv_note_run.font.italic = True
            cv_note_run.font.color.rgb = RGBColor(128, 128, 128)
            
            doc.add_paragraph()
            
            # Render CV content with formatting
            lines = cv_text.split('\n')
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue
                
                # Detect section headers
                is_section_header = (
                    (stripped.isupper() and len(stripped) > 3) or
                    (stripped.endswith(':') and len(stripped) < 60 and not stripped[0].islower()) or
                    stripped.startswith('#')
                )
                
                # Detect bullets
                is_bullet = stripped.startswith(('•', '-', '·', '*')) or re.match(r'^\d+[\.\)]\s+', stripped)
                
                # Clean markdown
                clean_text = re.sub(r'^#{1,3}\s+', '', stripped)
                clean_text = clean_text.replace('**', '').replace('__', '')
                
                if is_section_header:
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    run = p.add_run(clean_text.replace(':', ''))
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 51, 153)
                elif is_bullet:
                    bullet_text = re.sub(r'^[\s]*[\-\•\·\*]\s*', '', clean_text)
                    bullet_text = re.sub(r'^\d+[\.\)]\s*', '', bullet_text)
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(bullet_text).font.size = Pt(10)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(clean_text)
                    run.font.size = Pt(10)
            
            doc.add_paragraph()
        else:
            no_cv = doc.add_paragraph()
            no_cv.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = no_cv.add_run("CV generation was not available for this range.")
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("This report was auto-generated by Job Agent.")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.save(str(filepath))
    
    def export_all_jobs(self, jobs: List[Job], profile_name: str = "", filename: str = "job_listings.docx"):
        """Export all jobs to a single Word document (backward compatibility)."""
        doc = Document()
        
        title = doc.add_heading('Job Listings Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()
        
        summary = doc.add_paragraph()
        run = summary.add_run(f"Total Jobs Found: {len(jobs)}")
        run.font.bold = True
        run.font.size = Pt(12)
        
        doc.add_paragraph()
        
        by_platform = {}
        for job in jobs:
            platform_name = job.platform.value.upper() if hasattr(job.platform, 'value') else str(job.platform)
            if platform_name not in by_platform:
                by_platform[platform_name] = []
            by_platform[platform_name].append(job)
        
        for platform_name, platform_jobs in by_platform.items():
            heading = doc.add_heading(platform_name, level=1)
            heading.style.font.size = Pt(12)
            heading.style.font.color.rgb = RGBColor(80, 80, 80)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            header_cells = table.rows[0].cells
            header_cells[0].text = "Job Role"
            header_cells[1].text = "Company"
            header_cells[2].text = "Link"
            self._style_header_row(table.rows[0])
            
            for job in platform_jobs:
                row = table.add_row()
                row.cells[0].text = job.title
                row.cells[1].text = job.company
                
                link_para = row.cells[2].paragraphs[0]
                if job.url:
                    self._add_hyperlink(link_para, job.url, "Apply Link")
                else:
                    link_para.text = "N/A"
                
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(10)
            
            doc.add_paragraph()
        
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("This report was auto-generated by Job Agent.")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.save(str(self.output_dir / filename))


def export_jobs_to_word(jobs: List[Job], profile_name: str = "", filepath: str = None) -> str:
    """Original function for backward compatibility."""
    exporter = WordExporter()
    if filepath:
        exporter.export_all_jobs(jobs, profile_name, Path(filepath).name)
        return filepath
    exporter.export_all_jobs(jobs, profile_name)
    return str(Path("job_listings.docx").absolute())


def export_scored_jobs_to_word(scored_jobs: List[ScoredJob], profile_name: str = "",
                               output_dir: str = ".", cv_texts: Optional[Dict[str, str]] = None
                               ) -> List[Tuple[str, str, int, str]]:
    """
    Export jobs grouped by score range to separate Word files + PDF CVs.
    Returns list of (range_label, word_filepath, job_count, pdf_filepath) tuples.
    """
    exporter = WordExporter(output_dir)
    return exporter.export_jobs_with_scores(scored_jobs, profile_name, cv_texts)
